from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "测风数据邮件日报监测工具_客户使用手册.md"
OUTPUT = ROOT / "docs" / "测风数据邮件日报监测工具_客户使用手册.docx"


def set_font(run, size: float = 10.5, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def add_inline_text(paragraph, text: str, *, size: float = 10.5) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size)
            run.font.color.rgb = RGBColor(32, 84, 147)
        else:
            run = paragraph.add_run(part)
            set_font(run, size)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row):
            rows.append(row)
        index += 1
    return rows, index


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    title_done = False
    in_code = False
    code_lines: list[str] = []
    index = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.6)
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.add_run("\n".join(code_lines))
                set_font(run, 9)
                run.font.name = "Consolas"
                in_code = False
                code_lines = []
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            rows, index = parse_table(lines, index)
            if rows:
                column_count = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=column_count)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for row_index, row in enumerate(rows):
                    for column_index, value in enumerate(row):
                        cell = table.cell(row_index, column_index)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        paragraph = cell.paragraphs[0]
                        add_inline_text(paragraph, value, size=9.5)
                        if row_index == 0:
                            set_cell_shading(cell, "D9EAF7")
                            for run in paragraph.runs:
                                run.bold = True
                document.add_paragraph()
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1 and not title_done:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(18)
                run = paragraph.add_run(text)
                set_font(run, 20, bold=True)
                run.font.color.rgb = RGBColor(20, 65, 125)
                title_done = True
            else:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(10 if level == 2 else 6)
                paragraph.paragraph_format.space_after = Pt(5)
                run = paragraph.add_run(text)
                set_font(run, {2: 15, 3: 12.5, 4: 11}.get(level, 11), bold=True)
                run.font.color.rgb = RGBColor(20, 65, 125) if level <= 2 else RGBColor(31, 31, 31)
            index += 1
            continue

        list_match = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered_match = re.match(r"^\d+[.]\s+(.+)$", stripped)
        if list_match or numbered_match:
            text = (list_match or numbered_match).group(1)
            style = "List Bullet" if list_match else "List Number"
            paragraph = document.add_paragraph(style=style)
            add_inline_text(paragraph, text)
            index += 1
            continue

        quote_match = re.match(r"^>\s?(.*)$", stripped)
        if quote_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.right_indent = Cm(0.3)
            run = paragraph.add_run(quote_match.group(1))
            set_font(run, 10)
            run.font.color.rgb = RGBColor(89, 89, 89)
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline_text(paragraph, stripped)
        index += 1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("测风数据邮件日报监测工具使用手册")
    set_font(run, 8.5)
    run.font.color.rgb = RGBColor(127, 127, 127)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
